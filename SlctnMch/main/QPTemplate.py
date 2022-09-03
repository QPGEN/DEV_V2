# def fun(partc)
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches
from docx.enum.text import WD_BREAK_TYPE
from .models import Details
#import db elements
import os

def Data(H,CObj,COutcomes,partA,PartB,partC): #to create template
        sublist=Details.objects.filter(subcode=H[5]).values()
        subname=sublist[0]["subname"] 
        d= Document() #document docx object
    #--------------overall margin--------------------------------------
        sections=d.sections #sections is used here to set the margin size
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
#HEADER-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Header------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
   # H[i], where i =
   # 0 - "Internal Assessment",1 - IA- "1, 2 or 3", 2="date" , 3- "month", 4-"full date", 5- subjcode hc,
   # 6 - "branch" , 7 - Set, 8 - 50m or 100m, 9 - stud year, 10 - sem
        #H[2]->is not used, it contains only the date
        tdate=H[4].split('/')
        table=d.add_table(5,4)
        table.style="Table Grid"
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for i in range(5):
            for j in range(4):
                table.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1=table.cell(0,1)
        c3=table.cell(0,3)
        c10=table.cell(1,0)
        c12=table.cell(1,2)
        c13=table.cell(1,3)
        table.cell(0,0).merge(c1)
        table.cell(0,2).merge(c3)
        table.cell(1,1).merge(c12)
        table.cell(1,1).merge(c13)
     #   paragraph = c10.paragraphs[0]
     #   run = paragraph.add_run()
     #  run.add_picture('CIT Logo Simple - PNG.png')
        p11=table.cell(1,1).add_paragraph(text="CHENNAI INSTITUTE OF TECHNOLOGY\n Sarathy Nagar, Pudupedu, Chennai - 600 069."+"\n"+str(H[0])+"-"+str(H[1])+" "+str(H[3])+" "+str(tdate[2]))
        p11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p11=table.cell(1,1).add_paragraph(text=" ")
        table.cell(0,0).text="                                                                                                                      Reg. No.:"
        table.cell(2,0).text="Date /Time"
        table.cell(2,1).text=str(H[4])
        table.cell(2,2).text="Max. Marks"
        table.cell(2,3).text=str(H[8])
        table.cell(3,0).text="Subject with Code"
        table.cell(3,1).text=str(H[5])+ "-" + str(subname) # hc H6 to a name for now and use db later
        table.cell(3,2).text="Time"
        if H[8]=="50 Marks":
          table.cell(3,3).text="1.30 Hours"
        if H[8]=="100 Marks":
          table.cell(3,3).text="3 Hours" 
        table.cell(4,0).text="Branch"   #str(H[7])
        table.cell(4,1).text=str(H[6])+" - SET "+str(H[7])
        table.cell(4,2).text="Year/Semester"
        table.cell(4,3).text=str(H[9])+"/"+str(H[10])
        table.cell(2,1).width = Inches(5.0)
#COURSE OBJECTIVES--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Course Objectives-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("Course Objectives\nThe Student should be able ")

         #course objective size is a variable
        tableCObj=d.add_table(len(CObj)+1,2)
        # for _ in range(cobj+1):
        #     CourseObjectives.append([])
        S_no=tableCObj.cell(0,0).add_paragraph(text="Sl.no")
        S_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        Course_Objectives_Title=tableCObj.cell(0,1).add_paragraph(text="Course Objectives\n")
        Course_Objectives_Title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i in range(1,len(CObj)+1):
                tableCObj.cell(i,0).text="               "+str(i)
                tableCObj.cell(i,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i in range(1,len(CObj)+1):
                tableCObj.cell(i,1).text=CObj[i-1]

          
        tableCObj.cell(0,1).width=Inches(22.0)
        #table.cell(0,0).width=Inches(5)
        tableCObj.style="Table Grid"
#COURSE OUTCOMES---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Course Outcomes----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("Course Outcomes:\nOn Completion of the course the students will be able  ")

        tableCOut=d.add_table(len(COutcomes[0])+1,2)

        C_no=tableCOut.cell(0,0).add_paragraph(text="CO No.")
        C_no.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        Course_Outcomes_Title=tableCOut.cell(0,1).add_paragraph(text="Course Outcomes\n")
        Course_Outcomes_Title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i in range(1,len(COutcomes[0])+1):
              tableCOut.cell(i,0).text="            "+COutcomes[0][i-1]
              tableCOut.cell(i,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i in range(1,len(COutcomes[0])+1):
              tableCOut.cell(i,1).text=COutcomes[1][i-1]
        tableCOut.cell(0,1).width=Inches(20.0)

        #table.cell(0,0).width=Inches(5)
        tableCOut.style="Table Grid"
#PART A ----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------PartA--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("\n")

        table1a = d.add_table(1,1)
        table1a.style="Table Grid"
        for i in range(1):
            for j in range(1):
                table1a.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER   

        #INSERTING AND STYLING CONTENTS TO TABLE 1
        p1=table1a.cell(0,0).add_paragraph(text="BLOOMS TAXONOMY(BT)\nK1-Remembering , K2-Understanding, K3-Applying, K4-Analyzing, K5-Evaluating ,K6-Creating\nN/D-November/December A/M-April/June\n")
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p1bold=p1.add_run()
        # p1bold.bold=True
        #CREATING AND STYLING TABLE 2
        table2a = d.add_table(1,4)
        table2a.style = "Table Grid"
        table2a.cell(0,0).width = Inches(20.0)

        #INSERTING AND STYLING CONTENTS TO TABLE 2
        if H[8]=="50 Marks":
            p2=table2a.cell(0,0).add_paragraph(text="Part A-(5x2=10 marks)\n (Answer all the questions)\n")
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif H[8]=="100 Marks":
            p2=table2a.cell(0,0).add_paragraph(text="Part A-(10x2=20 marks)\n (Answer all the questions)\n")
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        p3=table2a.cell(0,1).add_paragraph(text="CO")
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p4=table2a.cell(0,2).add_paragraph(text="BT level")
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p5=table2a.cell(0,3).add_paragraph(text="Univ Qp \n reference")
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #CREATING AND STYLING TABLE 3
        if H[8]=="50 Marks":
           table3a = d.add_table(5,5)
           table3a.style = "Table Grid"
        elif H[8]=="100 Marks":
            table3a = d.add_table(10,5)
            table3a.style = "Table Grid"

        #INSERTING AND STYLING CONTENTS TO TABLE 3
        table3a.cell(0,1).width = Inches(15.0)
        qns=0
        if H[8]=="50 Marks":
             qns=5
        elif H[8]=="100 Marks":
             qns=10
        for i in range(qns):
            table3a.cell(i,0).add_paragraph(text="{}".format(i+1)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(qns):
            c=table3a.cell(i,1)
            table3a.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            list1=partA[0][i].split("~~")
            pic = c.paragraphs[0]
            run =pic.add_run()
            for k in list1:
                
                 if len(k)==0:
                     pass

                 elif k[0]!="$":
                     run.add_text(k)
                     run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                     
                 else: 
                     run.add_picture(k[1:])
                     run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
        for i in range(qns):
            table3a.cell(i,2).add_paragraph(text=partA[1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #CO
        for i in range(qns):
            table3a.cell(i,3).add_paragraph(text=partA[2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #BT Level
        for i in range(qns):
            table3a.cell(i,4).add_paragraph(text=partA[3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #Univ QP reference
#PART B --------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------PartB----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        d.add_paragraph("\n")
        roman=["i","ii","iii","iv","v","vi","vii","viii","ix","x"]
        # table1b = d.add_table(1,1)
        # table1b.style="Table Grid"
        # for i in range(1):
        #     for j in range(1):
        #         table1b.cell(i,j).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER   

        #INSERTING AND STYLING CONTENTS TO TABLE 1
        # p1=table1b.cell(0,0).add_paragraph(text="BLOOMS TAXONOMY(BT)\nK1-Remembering , K2-Understanding, K3-Applying, K4-Analyzing, K5-Evaluating ,K6-Creating\nN/D-November/December A/M-April/June\n")
        # p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       
        #CREATING AND STYLING TABLE 2
        table2b = d.add_table(1,5)
        table2b.style = "Table Grid"
        table2b.cell(0,0).width = Inches(20.0)

        #INSERTING AND STYLING CONTENTS TO TABLE 2
        if H[8]=="50 Marks":
              if partC[0]==1:
                    p2=table2b.cell(0,0).add_paragraph(text="Part B (2*12=24 marks)\n (Answer all the questions)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
              elif partC[0]==0:
                    p2=table2b.cell(0,0).add_paragraph(text="Part B (1*8 + 2*16=40 marks)\n (Answer all the questions)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif H[8]=="100 Marks":
              if partC[0]==1:
                    p2=table2b.cell(0,0).add_paragraph(text="Part B (5*13=65 marks)\n (Answer all the questions)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
              elif partC[0]==0:
                    p2=table2b.cell(0,0).add_paragraph(text="Part B (5*16=80 marks)\n (Answer all the questions)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p3=table2b.cell(0,1).add_paragraph(text="CO")
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p4=table2b.cell(0,2).add_paragraph(text="BT level\n")
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p5=table2b.cell(0,3).add_paragraph(text="Univ Qp \n reference")
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p6=table2b.cell(0,4).add_paragraph(text="Marks\nAlloted")
        p6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if H[8]=="50 Marks":
            #Part b=[6a,6b,7a,7b], 6a=[Qns,CO,BT_Level,Univ_QP_Reference,mark_alloted]
             qn_no=["6(a)","6(b)","7(a)","7(b)"]
             k=0
             for j in range(4):
                        
                        
                        if j%2!=0:
                            tableOR = d.add_table(1,1)
                            tableOR.style = "Table Grid"
                            tableOR.cell(0,0).text="                                                                                          OR"
                            tableOR.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            tableb = d.add_table(len(PartB[j][0]),6)
                            tableb.style = "Table Grid"
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for i in range(len(PartB[j][0])):
                                        c=tableb.cell(i,1)
                                        tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=PartB[j][0][i].split("~~")

                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if j!=3:
                             tableBlank = d.add_table(1,1)
                             tableBlank.style = "Table Grid"
                            
                        else:

                                tableb = d.add_table(len(PartB[j][0]),6)
                                tableb.style = "Table Grid"
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for i in range(len(PartB[j][0])):
                                        c=tableb.cell(i,1)
                                        tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=PartB[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                                    #questions
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        k=k+1
                        
             if partC[0]==0:
                    
                    for j in range(4,6):
                        if j%2!=0:
                            tableOR = d.add_table(1,1)
                            tableOR.style = "Table Grid"
                            tableOR.cell(0,0).text="                                                                                          OR"
                            tableOR.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            tableb = d.add_table(len(PartB[j][0]),6)
                            tableb.style = "Table Grid"
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format("8(b)"+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for i in range(len(PartB[j][0])):
                                        c=tableb.cell(i,1)
                                        tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=PartB[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                                tableb = d.add_table(len(PartB[j][0]),6)
                                tableb.style = "Table Grid"
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format("8(a)"+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,1).text="\n"+PartB[j][0][i]#+"\n"
                                    tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    #questions
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        elif H[8]=="100 Marks":
             qn_no=["11(a)","11(b)","12(a)","12(b)","13(a)","13(b)","14(a)","14(b)","15(a)","15(b)"]
             k=0
             for j in range(10):
                        
                        
                        if j%2!=0:
                            tableOR = d.add_table(1,1)
                            tableOR.style = "Table Grid"
                            tableOR.cell(0,0).text="                                                                                          OR"
                            tableOR.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            tableb = d.add_table(len(PartB[j][0]),6)
                            tableb.style = "Table Grid"
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for i in range(len(PartB[j][0])):
                                        c=tableb.cell(i,1)
                                        tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=PartB[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                            for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if j!=9:
                             tableBlank = d.add_table(1,1)
                             tableBlank.style = "Table Grid"
                            
                        else:

                                tableb = d.add_table(len(PartB[j][0]),6)
                                tableb.style = "Table Grid"
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for i in range(len(PartB[j][0])):
                                        c=tableb.cell(i,1)
                                        tableb.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=PartB[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,2).add_paragraph(text=PartB[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #CO
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,3).add_paragraph(text=PartB[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #BT Level
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,4).add_paragraph(text=PartB[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    #Marks Alloted
                                for i in range(len(PartB[j][0])):
                                    tableb.cell(i,5).add_paragraph(text=PartB[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        k=k+1


             
#PART C----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------Part C---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        if partC[0]==1:
                d.add_paragraph("\n")
                roman=["i","ii","iii","iv","v","vi","vii","viii","ix","x"]
                ab=["b","a"]

                #CREATING AND STYLING TABLE 2
                table2c = d.add_table(1,5)
                table2c.style = "Table Grid"
                table2c.cell(0,0).width = Inches(20.0)
                
                #INSERTING AND STYLING CONTENTS TO TABLE 2
                if H[8]=="50 Marks":
                    p2=table2c.cell(0,0).add_paragraph(text="Part C (1*16=16 marks)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif H[8]=="100 Marks":
                    p2=table2c.cell(0,0).add_paragraph(text="Part C (1*15=15 marks)\n")
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                
                p3=table2c.cell(0,1).add_paragraph(text="CO")
                p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p4=table2c.cell(0,2).add_paragraph(text="BT level\n")
                p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p5=table2c.cell(0,3).add_paragraph(text="Univ Qp \n reference")
                p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p6=table2c.cell(0,4).add_paragraph(text="Marks\nAlloted")
                p6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER       
         
                qn_no=["8.","16."] 
                k=-1
                if H[8]=="50 Marks":
                    k=0
                elif H[8]=="100 Marks":
                    k=1
                for j in range(1,3):
                            if j%2==0:
                                tableOR = d.add_table(1,1)
                                tableOR.style = "Table Grid"
                                tableOR.cell(0,0).text="                                                                                          OR"
                                tableOR.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                tablec = d.add_table(len(partC[j][0]),6)
                                tablec.style = "Table Grid"
                                for i in range(len(partC[j][0])):
                                        tablec.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+ab[j%2]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for i in range(len(partC[j][0])):
                                        c=tablec.cell(i,1)
                                        tablec.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=partC[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    #questions
                                for i in range(len(partC[j][0])):
                                        tablec.cell(i,2).add_paragraph(text=partC[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #CO
                                for i in range(len(partC[j][0])):
                                        tablec.cell(i,3).add_paragraph(text=partC[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #BT Level
                                for i in range(len(partC[j][0])):
                                        tablec.cell(i,4).add_paragraph(text=partC[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #Marks Alloted
                                for i in range(len(partC[j][0])):
                                        tablec.cell(i,5).add_paragraph(text=partC[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                    tablec = d.add_table(len(partC[j][0]),6)
                                    tablec.style = "Table Grid"
                                    for i in range(len(partC[j][0])):
                                        tablec.cell(i,0).add_paragraph(text="{}".format(qn_no[k]+ab[j%2]+"("+roman[i]+")")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    for i in range(len(partC[j][0])):
                                        c=tablec.cell(i,1)
                                        tablec.cell(i,1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                        list1=partC[j][0][i].split("~~")
                                        pic = c.paragraphs[0]
                                        run =pic.add_run()
                                        for t in list1:
                                            
                                            if len(t)==0:
                                                pass

                                            elif t[0]!="$":
                                                run.add_text(t)
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                                
                                            else: 
                                                run.add_picture(t[1:])
                                                run.add_break(WD_BREAK_TYPE.TEXT_WRAPPING)
                                    for i in range(len(partC[j][0])):
                                        tablec.cell(i,2).add_paragraph(text=partC[j][1][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #CO
                                    for i in range(len(partC[j][0])):
                                        tablec.cell(i,3).add_paragraph(text=partC[j][2][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #BT Level
                                    for i in range(len(partC[j][0])):
                                        tablec.cell(i,4).add_paragraph(text=partC[j][3][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        #Marks Alloted
                                    for i in range(len(partC[j][0])):
                                        tablec.cell(i,5).add_paragraph(text=partC[j][4][i]).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       
        d.add_paragraph("\n\n\n")
        d.add_paragraph("Prepared by                                                                 Verified by                                                                 Approved by").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        d.save(H[5]+" "+H[7]+".docx")
   
